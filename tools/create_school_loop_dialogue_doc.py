from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"C:\Users\33574\Documents\剧本杀\docs\学校循环台词第一版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    for name, size in [("Normal", 10.5), ("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        if name.startswith("Heading"):
            style.font.bold = True


def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "2F3A4A")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if row_index % 2 == 0:
                set_cell_shading(cells[i], "F4F6F8")
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


stage_lines = [
    ["阶段1 校门口", "沈晴", "所有人先别堵在门口。迟到的名字我会记下来，想解释的等到教室再说。", "控场过渡"],
    ["阶段1 校门口", "林澈", "许骁，今天别拦我。我不想因为你再被老师点名。", "被许骁拦住"],
    ["阶段1 校门口", "许骁", "急什么？学霸也有怕迟到的时候？把你手里的东西给我看看。", "制造冲突"],
    ["阶段1 校门口", "唐玥", "你们吵够了吗？我还有东西要取，别挡路。", "提前埋手机线"],
    ["阶段1 校门口", "何岚", "别碰我的画夹，里面夹的是今天要交的材料。", "可见线索为画夹门禁贴纸"],
    ["阶段1 走廊", "周燃", "器材室钥匙又被收走了。等会儿谁要搬器材，别临时找我。", "体育生线入口"],
    ["阶段1 结束", "沈晴", "铃声快响了。所有人进教室，顾老师已经在点名本旁边等着了。", "进入阶段2"],
    ["阶段2 教室", "顾老师", "今天的点名照常。沈晴，你记录迟到。其他人把无关的东西收起来。", "控场"],
    ["阶段2 教室", "沈晴", "魏眠，你又坐错位置了。算了，先别动，我会改记录。", "胆小鬼线入口"],
    ["阶段2 教室", "魏眠", "我没有坐错。我只是觉得那个位置今天不太对。", "暗示异常"],
    ["阶段2 教室", "林澈", "题目可以换，答案也可以换，但人不会突然变聪明。你要证明，就回答我。", "学霸问答入口"],
    ["阶段2 结束", "顾老师", "课到这里。下一阶段去操场，周燃负责带队。", "进入阶段3"],
    ["阶段3 操场", "周燃", "别站着看。要钥匙可以，先比一局。赢了我就让你帮忙。", "扳手腕入口"],
    ["阶段3 操场", "魏眠", "我不是故意看到的。那个人从功能教室后门进去，又从另一边出来。", "目击线"],
    ["阶段3 结束", "周燃", "体育课结束。去食堂，别把器材留在场地上。", "进入阶段4"],
    ["阶段4 食堂", "许骁", "那个柜子不是你的。你要翻，就看你能不能比我快。", "抢柜子入口"],
    ["阶段4 食堂", "唐玥", "手机还给我。你拿着它也没用，里面的记录你看不懂。", "富家女线入口"],
    ["阶段4 结束", "沈晴", "功能教室预约开始了。需要材料的人现在过去，晚了就锁门。", "进入阶段5"],
    ["阶段5 美术室", "何岚", "配方不是背出来的，是手跟得上才做得出来。你要用，就按我的节奏来。", "节奏对抗入口"],
    ["阶段5 广播室", "林澈", "密码只开一次门。你要是用错，广播室今天就别再想进。", "广播室限制"],
    ["阶段5 结束", "顾老师", "功能教室关闭。所有人到办公室外集合，今天的问题该收尾了。", "进入阶段6"],
    ["阶段6 办公室", "顾老师", "你今天拿到了很多东西。但我想知道，你到底记住了什么。", "最终线入口"],
    ["阶段6 普通结算", "顾老师", "还不够。你只是改变了几件事，还没有理解这一天为什么重复。", "失败循环"],
    ["阶段6 真结局", "顾老师", "如果你连代价都能说清楚，那这一天就没有理由再困住你。", "成功"],
]


route_lines = [
    ["学霸 林澈", "初次信任", "你帮我不是因为好心。你一定想从我这里拿什么。", "选项：我想知道广播室密码。/ 我只是看不惯他。"],
    ["学霸 林澈", "知识问答开始", "不用小游戏。你自己设计问题，主角通过普通对话回答。这里建议放三题，每题一个选择。", "答对三题进入成功节点，答错任意关键题进入警觉节点。"],
    ["学霸 林澈", "成功", "行，你不是乱猜的。广播室密码我只说一次，记住了就别再问。", "给予密码或可保留线索。"],
    ["学霸 林澈", "失败", "你根本不知道答案。别再靠近我的桌子。", "本轮学霸线锁死。"],
    ["体育生 周燃", "挑战", "想拿器材室钥匙？先让我看看你有没有资格碰那些东西。", "触发扳手腕。"],
    ["体育生 周燃", "胜利", "可以。你力气够，钥匙给你十分钟。别让我发现你乱用。", "获得当天钥匙。"],
    ["体育生 周燃", "失败", "不行。你连我手都压不住，进去只会添乱。", "本轮体育生线暂停。"],
    ["小混混 许骁", "威胁", "你知道太多，就会变成麻烦。麻烦通常会自己消失。", "可选择硬顶或绕开。"],
    ["小混混 许骁", "柜子争夺", "我数三下。谁先找到，东西归谁。", "触发抢柜子。"],
    ["小混混 许骁", "失败", "你慢了。下次想翻我的柜子，先学会别让人看见。", "本轮转移藏品。"],
    ["班长 沈晴", "点名表", "记录不是为了告状，是为了知道谁在什么时候不该出现。", "触发记忆翻牌。"],
    ["班长 沈晴", "成功", "这份路线表你别拿走，看一眼就够了。被老师发现我也解释不了。", "获得巡逻路线。"],
    ["班长 沈晴", "警觉", "你问得太细了。今天的记录我自己整理。", "本轮班长线锁死。"],
    ["富家女 唐玥", "手机", "你想看我的手机？可以，先证明你能在他们之前拿到它。", "触发抢柜子或包内争夺。"],
    ["富家女 唐玥", "成功", "别乱翻相册。你要的只有那张监控截图。", "获得监控盲区。"],
    ["胆小鬼 魏眠", "安抚", "别在他们面前问我。我说了，他们都会知道是我说的。", "私下对话。"],
    ["胆小鬼 魏眠", "记忆", "我只记得几个位置。你能按顺序找出来，我就继续说。", "触发记忆翻牌。"],
    ["胆小鬼 魏眠", "失败", "不对，不是那里。算了，我不能再说了。", "扣分并换组，分数不足则失败。"],
    ["艺术生 何岚", "画夹", "你看画可以，但别翻夹层。那不是给你看的。", "门禁贴纸为可见线索。"],
    ["艺术生 何岚", "节奏", "颜料比例错一点，成品就不像。跟上，我只做一遍。", "触发节奏对抗。"],
    ["艺术生 何岚", "成功", "这个伪装只能骗过一次。用完就丢，别带回我这里。", "获得当天伪装道具。"],
]


teacher_questions = [
    ["问题1", "顾老师", "谁是今天最先偏离原本行动的人？", "正确答案由你设计，可绑定前几轮线索。"],
    ["问题2", "顾老师", "循环真正开始的地点在哪里？", "建议答案：校门口或办公室外，按地图机关决定。"],
    ["问题3", "顾老师", "为什么只有主角能记住之前的循环？", "建议答案：主角携带或触碰过某个保留道具。"],
]


def build_doc():
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("学校循环台词第一版")
    add_p(doc, "本文档整理第一版可用台词，供 NPC 玩家控场和后续 JSON 对话制作使用。学霸线的核心玩法已改为普通对话知识问答，题目内容保留给地图作者自行设计。")
    add_p(doc, "台词尽量使用 Minecraft 可以表现的内容，例如可见物品、站位、钥匙、门禁贴纸、柜子、广播、点名表和玩家移动。没有使用气味、镜头特写或心理描写作为关键线索。")

    doc.add_heading("一 使用规则", level=1)
    add_table(doc, ["规则", "说明"], [
        ["NPC 控场", "每阶段至少一个 NPC 负责公开宣布集合、移动或场景结束。"],
        ["自由说话", "NPC 可以临场发挥，但不能直接透露正确路线和最终答案。"],
        ["硬规则", "信物、体力、阶段推进、小游戏胜负和道具发放以模组或管理员判定为准。"],
        ["学霸问答", "使用普通对话选择实现，不再使用节奏对抗。"],
    ], [4.0, 12.0])

    doc.add_heading("二 阶段控场台词", level=1)
    add_table(doc, ["位置", "说话人", "台词", "用途"], stage_lines, [3.0, 2.4, 8.5, 3.0])

    doc.add_heading("三 角色路线台词", level=1)
    add_table(doc, ["角色", "节点", "台词", "玩法或结果"], route_lines, [2.6, 3.0, 8.0, 4.0])

    doc.add_heading("四 学霸知识问答模板", level=1)
    add_p(doc, "学霸线不再使用小游戏。下面是普通对话节点模板，你可以把题目替换成自己的知识题、地图谜题或剧情题。")
    add_table(doc, ["节点", "林澈台词", "回答选项", "跳转"], [
        ["题目一", "第一题。你说你知道答案，那就别犹豫。问题一由作者填写。", "A 正确答案 / B 干扰项 / C 干扰项", "正确进入题目二，错误进入失败"],
        ["题目二", "第二题。上一题可能是巧合，这一题不会。问题二由作者填写。", "A 干扰项 / B 正确答案 / C 干扰项", "正确进入题目三，错误进入失败"],
        ["题目三", "最后一题。答出来，我就相信你不是在蒙。问题三由作者填写。", "A 干扰项 / B 干扰项 / C 正确答案", "正确进入成功，错误进入失败"],
        ["成功", "行，你不是乱猜的。广播室密码我只说一次。", "继续", "获得密码线索"],
        ["失败", "你根本不知道答案。别再靠近我的桌子。", "结束", "本轮学霸线锁死"],
    ], [2.4, 6.2, 5.0, 4.0])

    doc.add_heading("五 老师最终问答", level=1)
    add_table(doc, ["题号", "说话人", "台词", "设计说明"], teacher_questions, [2.0, 2.4, 8.0, 5.0])

    doc.add_heading("六 NPC 过渡提醒", level=1)
    add_table(doc, ["触发条件", "推荐台词", "作用"], [
        ["主角体力即将用完", "你看起来撑不住了。把这句话说完，我们就该去下一个地方。", "解释体力耗尽后对话结束才切阶段"],
        ["玩家拖延不移动", "下一处已经开放了，留在这里也不会有新东西。", "防止场景卡住"],
        ["玩家问无法实现的感官线索", "看得到的东西才算证据。你要找，就看桌面、柜子和门口。", "把线索拉回可交互物品"],
        ["普通循环失败", "今天结束。你带不走那些信物，但你应该记住自己哪里慢了。", "强调信物不跨循环"],
        ["真结局前", "如果你真能一次说清所有人的结局，就进来。否则，明天还会一样。", "引导老师线"],
    ], [4.0, 8.0, 4.0])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
