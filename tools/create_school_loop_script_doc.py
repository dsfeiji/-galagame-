from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"C:\Users\33574\Documents\剧本杀\docs\学校循环剧本第一版修订.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        if "Heading" in name:
            style.font.bold = True


def add_p(doc, text="", bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_lead:
        run = paragraph.add_run(bold_lead)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr[index], header, True)
        set_cell_shading(hdr[index], "2F3A4A")
        for run in hdr[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[index], "F4F6F8")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


characters = [
    ["主角", "protagonist", "唯一能推动阶段的人。保留记忆，负责规划当天路线。", "30 点阶段体力预算的实际使用者。"],
    ["学霸 林澈", "student_scholar", "掌握试卷答案、广播室密码和教师电脑口令。", "节奏对抗"],
    ["体育生 周燃", "student_athlete", "掌握器材室钥匙，能搬动重物。", "扳手腕"],
    ["小混混 许骁", "student_bully", "藏有违禁物和威胁证据。", "抢柜子"],
    ["班长 沈晴", "student_monitor", "掌握点名表、老师巡逻表和迟到记录。", "记忆翻牌"],
    ["富家女 唐玥", "student_rich", "掌握手机、监控盲区和交易记录。", "抢柜子"],
    ["胆小鬼 魏眠", "student_coward", "知道谁在什么时候进入过关键房间。", "记忆翻牌"],
    ["艺术生 何岚", "student_artist", "能制作伪装道具，接触颜料和清洁剂。", "节奏对抗"],
    ["老师 顾老师", "teacher", "知道循环存在，是最终信物持有者。", "最终审问和同轮全清检查"],
]


phases = [
    ["1", "校门口和走廊", "入校、观察当天状态、拿初始道具。", "5", "选择第一条主线入口，错误搭话会浪费体力。"],
    ["2", "教室", "早读、点名、学霸和班长信息交叉。", "5", "拿到密码或巡逻表，决定当天是否能全收集。"],
    ["3", "操场", "体育课、公开冲突、器材室钥匙。", "5", "体育生线和胆小鬼线开始互相影响。"],
    ["4", "食堂和小卖部", "交易、手机、违禁物、学生关系曝光。", "5", "富家女和小混混线关键阶段。"],
    ["5", "功能教室区", "器材室、美术室、广播室、实验准备间。", "5", "集中执行大部分意外，体力安排最紧。"],
    ["6", "办公室和放学后", "老师线、最终检查、本轮结算。", "5", "同轮 7 个学生信物齐全才可拿老师信物。"],
]

npc_control = [
    ["1", "校门口和走廊", "沈晴", "组织排队入校，提醒迟到记录开始。", "主角靠近校门或倒计时结束后，由沈晴提示所有人进入教室。"],
    ["2", "教室", "顾老师和沈晴", "顾老师点名，沈晴补充学生状态。", "点名结束后，顾老师宣布去操场上体育课，周燃带头离开。"],
    ["3", "操场", "周燃", "组织热身和器材搬运，给扳手腕挑战理由。", "体育课结束时，周燃提醒去食堂，小混混和富家女提前离开。"],
    ["4", "食堂和小卖部", "许骁和唐玥", "围绕交易、手机和柜子争夺制造冲突。", "食堂广播响起后，沈晴提醒功能教室预约开始。"],
    ["5", "功能教室区", "何岚和林澈", "何岚控制美术室材料，林澈控制广播室信息。", "功能教室关闭铃响后，顾老师要求所有人到办公室外集合。"],
    ["6", "办公室和放学后", "顾老师", "进行最终审问或普通结算。", "若条件不足，顾老师宣布今天结束，循环重开。若条件满足，进入真结局。"],
]


routes = [
    {
        "title": "学霸线 林澈",
        "token": "破损的满分试卷",
        "carry": "广播室密码可跨循环保留",
        "steps": [
            "阶段 1 在走廊帮林澈挡住小混混，消耗 1 体力，获得一次信任。",
            "阶段 2 与林澈进行节奏对抗，表现为抢答。胜利后获得试卷答案和广播室密码。",
            "阶段 5 用广播室密码播放错误集合铃，让林澈提前进入空教室。",
            "阶段 5 使用清洁剂制造地面湿滑，消耗 2 体力，林澈退场。",
            "结算获得信物。若阶段 2 输掉小游戏，林澈警觉，本轮无法再单独引走。",
        ],
    },
    {
        "title": "体育生线 周燃",
        "token": "断裂的护腕",
        "carry": "器材室钥匙复制信息可跨循环保留，但钥匙本身不保留",
        "steps": [
            "阶段 1 听到周燃抱怨器材室锁坏了，记录入口信息。",
            "阶段 3 在操场与周燃扳手腕，胜利后他承认主角够资格帮忙搬器材。",
            "阶段 3 获得器材室钥匙，消耗 1 体力。",
            "阶段 5 在器材室松动篮球架底座，消耗 2 体力。",
            "阶段 5 引导周燃检查篮球架，制造器材倒塌事故，获得信物。",
        ],
    },
    {
        "title": "小混混线 许骁",
        "token": "烧焦的纸条",
        "carry": "他藏东西的柜号可跨循环保留",
        "steps": [
            "阶段 1 被许骁勒索，可以交出普通物品避免冲突，也可以硬顶消耗 1 体力。",
            "阶段 4 右键储物柜触发抢柜子小游戏，主角胜利后拿到违禁物和威胁纸条。",
            "阶段 4 若点错柜子不会扣分，但会浪费时间，输掉则许骁转移物品。",
            "阶段 5 把违禁物放进实验准备间，并用广播引许骁过去。",
            "阶段 5 制造烟雾警报误触，许骁在混乱中退场，获得信物。",
        ],
    },
    {
        "title": "班长线 沈晴",
        "token": "缺角的点名牌",
        "carry": "老师巡逻路线可跨循环保留",
        "steps": [
            "阶段 2 帮沈晴整理点名表，消耗 1 体力。",
            "阶段 2 触发记忆翻牌，3x5 共 15 格中记住 6 个发光位置，复原老师巡逻表。",
            "阶段 2 成功后获得老师巡逻路线和迟到名单。",
            "阶段 5 利用巡逻空档修改功能教室预约记录，消耗 1 体力。",
            "阶段 6 指认沈晴伪造点名导致混乱，迫使她独自返回教室，被之前布置的机关退场，获得信物。",
        ],
    },
    {
        "title": "富家女线 唐玥",
        "token": "碎裂的手机挂坠",
        "carry": "监控盲区信息可跨循环保留",
        "steps": [
            "阶段 1 观察唐玥和许骁交易，记下手机型号。",
            "阶段 4 在小卖部附近触发抢柜子，表现为争夺手机和包内交易记录。",
            "阶段 4 胜利后获得监控盲区截图和一张购物小票。",
            "阶段 5 使用购物小票换取备用清洁剂，消耗 1 体力。",
            "阶段 6 用手机信息让唐玥离开办公室门口，制造楼梯间意外，获得信物。",
        ],
    },
    {
        "title": "胆小鬼线 魏眠",
        "token": "揉皱的目击便签",
        "carry": "他看到的房间顺序可跨循环保留",
        "steps": [
            "阶段 2 不要当众逼问魏眠，否则本轮锁线。",
            "阶段 3 在操场角落安抚魏眠，消耗 1 体力。",
            "阶段 3 触发记忆翻牌，复原他看到的 6 个关键位置。点错扣 1 分并进入下一组。",
            "阶段 3 成功后获得目击证词，知道谁能打开功能教室后门。",
            "阶段 5 利用证词诱导魏眠去确认后门，制造误锁事故，获得信物。",
        ],
    },
    {
        "title": "艺术生线 何岚",
        "token": "染色的画笔",
        "carry": "伪装配方可跨循环保留",
        "steps": [
            "阶段 1 夸赞何岚手里的画夹，但不能追问画夹中夹着的门禁贴纸，否则本轮警觉。",
            "阶段 5 在美术室触发节奏对抗，表现为快速调色和图案复刻。",
            "阶段 5 胜利后获得伪装道具和颜料配方。",
            "阶段 5 使用伪装道具进入实验准备间，消耗 1 体力。",
            "阶段 6 让何岚误以为作品被毁，引她返回美术室并触发先前布置，获得信物。",
        ],
    },
]


def build_doc():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("学校循环剧本第一版")
    add_p(doc, "本文档是 Minecraft 剧本杀小地图的第一版完整玩法剧本。它按照现有模组能力设计，包括角色认领、阶段统一推进、每阶段 5 点体力、对话分支、可保留道具、不可保留信物、以及扳手腕、抢柜子、节奏对抗、记忆翻牌四类小游戏。")
    add_p(doc, "核心目标是让玩家通过多次循环掌握路线，然后在某一轮内收集 7 个学生信物并触发老师最终线。信物不跨循环保留，部分线索和道具可以保留。")

    doc.add_heading("一 核心规则", level=1)
    add_bullets(doc, [
        "一轮循环分为 6 个阶段，一个主要场景对应一个阶段。",
        "每个阶段主角拥有 5 点体力，进入新阶段时体力恢复为 5。",
        "体力耗尽不会打断当前对话，当前对话关闭后再进入下一阶段。",
        "学生信物和老师信物只在当前循环有效，下一轮开始全部清空。",
        "密码、路线情报、配方、柜号等记忆类信息可以跨循环保留。",
        "主角是唯一能推动剧情阶段的人，其他玩家只能参与对话和小游戏。",
    ])

    doc.add_heading("二 角色总表", level=1)
    add_table(doc, ["角色", "角色 ID", "剧情功能", "核心小游戏"], characters, [3.0, 3.5, 7.0, 3.5])

    doc.add_heading("三 阶段结构", level=1)
    add_table(doc, ["阶段", "场景", "主要内容", "体力", "设计重点"], phases, [1.3, 3.0, 6.0, 1.3, 5.0])

    doc.add_heading("四 NPC 控场和过渡", level=1)
    add_p(doc, "本地图不是所有 NPC 都被动等待主角点击。部分场景需要由 NPC 玩家控场，让阶段切换自然发生。控场 NPC 的任务不是替主角做决定，而是用公开台词、站位移动和物品展示，把所有玩家带到下一个场景。")
    add_table(doc, ["阶段", "场景", "控场 NPC", "控场职责", "过渡方式"], npc_control, [1.2, 3.0, 3.0, 5.8, 6.2])

    doc.add_heading("五 循环和保留规则", level=1)
    add_p(doc, "每次进入新循环时，所有学生信物和老师信物清空。玩家不能通过多轮慢慢攒齐信物，最终通关必须在同一轮内完成。")
    add_p(doc, "可跨循环保留的内容主要是信息和少量记忆类道具，例如密码、柜号、路线表、配方、监控盲区。不可保留的内容包括学生信物、老师信物、当天钥匙、当天手机、当天违禁物和临时伪装道具。")
    add_table(doc, ["类型", "是否保留", "例子", "用途"], [
        ["信物", "不保留", "破损试卷、断裂护腕、手机挂坠", "最终线同轮检查"],
        ["当天物品", "不保留", "钥匙、手机、违禁物、清洁剂", "限制每轮行动顺序"],
        ["记忆信息", "保留", "密码、柜号、巡逻路线、配方", "帮助下一轮节省体力"],
        ["关键路线理解", "保留", "谁怕谁、谁能引开谁", "支撑全收集路线规划"],
    ], [3.0, 2.2, 5.0, 6.0])

    doc.add_heading("六 小游戏设计", level=1)
    add_table(doc, ["小游戏", "触发方式", "规则", "剧情用途"], [
        ["扳手腕", "对话或右键器材", "双方疯狂点击，进度条向对方方向移动，到边界结束。", "力量对抗和资格证明"],
        ["抢柜子", "右键柜子或对话争夺", "点对加分，点错不扣分，点对后换目标。", "争夺手机、证据和藏品"],
        ["节奏对抗", "对话或右键物品", "点击后立刻进入下一轮，绿色区位置和宽度随机变化。", "抢答、调色、快速操作"],
        ["记忆翻牌", "对话或右键表格", "3x5 共 15 格，每轮记住 6 个发光格子，点错扣 1 分并换组。", "复原路线、目击信息和巡逻表"],
    ], [2.5, 4.0, 7.0, 4.5])

    doc.add_heading("七 七条学生路线", level=1)
    for route in routes:
        doc.add_heading(route["title"], level=2)
        add_p(doc, route["token"], bold_lead="本轮信物：")
        add_p(doc, route["carry"], bold_lead="可保留内容：")
        add_bullets(doc, route["steps"])

    doc.add_heading("八 老师最终线", level=1)
    add_p(doc, "顾老师知道循环存在。她平时不会直接阻止主角，但会记录每天异常。只有在同一轮内获得 7 个学生信物，阶段 6 才会开启老师最终线。")
    add_bullets(doc, [
        "阶段 6 进入办公室前，系统检查本轮信物数量。",
        "若不足 7 个，老师只会说主角还没有理解这一天，然后进入普通循环结算。",
        "若正好拥有 7 个学生信物，老师进入最终审问。",
        "最终审问中，老师会询问三个问题：谁最先被改变、哪个地点是循环起点、主角为什么还能记得。",
        "三个问题可以由前几轮保留线索推理出来。答错不会死亡，但本轮失败并进入下一循环。",
        "答对后获得老师信物，触发真结局。",
    ])

    doc.add_heading("九 推荐全收集路线", level=1)
    add_table(doc, ["阶段", "行动顺序", "体力消耗", "关键前置"], [
        ["1", "帮林澈挡许骁；观察唐玥交易；与何岚建立信任；记录周燃抱怨。", "4", "最好已知道许骁柜号"],
        ["2", "赢林澈节奏对抗；赢沈晴记忆翻牌；安抚魏眠但不公开逼问。", "5", "需要上一轮知道老师抽屉无须再查"],
        ["3", "赢周燃扳手腕；拿器材室钥匙；赢魏眠记忆翻牌。", "5", "需要知道后门顺序"],
        ["4", "赢许骁抢柜子；赢唐玥抢柜子；换清洁剂。", "5", "需要柜号和手机线索"],
        ["5", "布置篮球架、广播引林澈、准备烟雾、制作伪装、处理后门。", "5", "必须精准顺序，否则体力不够"],
        ["6", "依次完成七个信物结算；进入办公室答老师三问。", "5", "本轮七个学生信物齐全"],
    ], [1.3, 9.0, 2.2, 5.0])

    doc.add_heading("十 阶段详细事件", level=1)
    stage_details = [
        ("阶段 1 校门口和走廊", [
            "主角醒在校门口，身上没有信物，只保留上一轮允许保留的信息。",
            "林澈被许骁拦住，主角可以选择帮忙或旁观。",
            "唐玥和许骁短暂交易，若主角靠近可看到手机挂坠。",
            "周燃抱怨器材室钥匙被老师收走。",
            "何岚经过走廊，手里拿着画夹，画夹角落贴着一张功能教室门禁贴纸。",
            "沈晴负责公开提醒所有人进入教室，她的台词用于把阶段 1 过渡到阶段 2。",
        ]),
        ("阶段 2 教室", [
            "顾老师点名，沈晴协助记录迟到。",
            "林澈可以被挑战抢答，胜利后给出广播室密码。",
            "沈晴会让主角帮忙整理点名表，触发记忆翻牌。",
            "魏眠紧张地看向功能教室方向，当众追问会锁线。",
            "顾老师宣布去操场，周燃在门口催促所有人移动到阶段 3。",
        ]),
        ("阶段 3 操场", [
            "体育课开始，周燃接受扳手腕挑战。",
            "主角赢后可以参与搬器材并短暂获得钥匙。",
            "魏眠在操场角落透露他看到有人从后门进出。",
            "记忆翻牌用于复原魏眠看到的 6 个位置。",
            "周燃以体育课结束为理由带队去食堂，同时许骁和唐玥先一步离开制造阶段 4 冲突。",
        ]),
        ("阶段 4 食堂和小卖部", [
            "许骁把违禁物藏在柜子里，唐玥试图拿回手机。",
            "抢柜子用于争夺藏品，点错不扣分但浪费时间。",
            "主角可用小票或信息换清洁剂。",
            "若没有上一轮信息，体力通常不够同时完成许骁和唐玥两线。",
            "沈晴听到广播后提醒功能教室预约开始，推动玩家进入阶段 5。",
        ]),
        ("阶段 5 功能教室区", [
            "这里是全收集最紧张的阶段。",
            "主角要使用器材室钥匙、广播室密码、清洁剂、伪装配方和后门信息。",
            "每个意外都消耗体力，错误顺序会导致无法在本轮完成全部信物。",
            "何岚的节奏对抗在美术室触发，成功后获得伪装道具。",
            "何岚和林澈分别控制美术室与广播室入口，负责让玩家相信这些房间有明确使用规则。",
        ]),
        ("阶段 6 办公室和放学后", [
            "本轮信物不足时进入普通结算。",
            "七个学生信物齐全时，老师最终线开启。",
            "老师不会被小游戏击败，而是通过对话和线索答案判断主角是否真正理解循环。",
            "答对三问后获得老师信物，循环结束。",
            "顾老师负责最终控场。没有达成条件时，她用结算台词把所有人带回下一轮。",
        ]),
    ]
    for heading, items in stage_details:
        doc.add_heading(heading, level=2)
        add_bullets(doc, items)

    doc.add_heading("十一 关键对话样稿", level=1)
    add_table(doc, ["场景", "说话人", "台词", "回答和结果"], [
        ["教室", "林澈", "你又知道这道题的答案？不可能，每天题都不一样。", "回答一：我只是记性好，进入节奏对抗。回答二：我看过答案，林澈警觉。"],
        ["操场", "周燃", "想拿钥匙？先让我看看你有没有这个力气。", "接受挑战消耗 1 体力，进入扳手腕。拒绝则本轮体育线暂停。"],
        ["储物柜", "许骁", "别碰那个柜子，里面不是你该看的东西。", "强行翻找进入抢柜子。若胜利获得违禁物。"],
        ["教室", "沈晴", "老师今天走哪条路线，我好像记混了。", "帮她整理触发记忆翻牌，成功获得巡逻路线。"],
        ["办公室", "顾老师", "你收集的是证据，还是你自己逃出去的理由？", "持有 7 个学生信物时进入最终三问。"],
    ], [2.2, 2.3, 7.3, 5.0])

    doc.add_heading("十二 NPC 玩家执行提示", level=1)
    add_bullets(doc, [
        "NPC 玩家可以自由说话，但不能直接告诉主角正确答案，只能根据本阶段掌握的信息给提示。",
        "每个阶段至少安排一个控场 NPC。控场 NPC 负责公开宣布移动、集合、点名或房间开放。",
        "所有重要线索必须能通过可见物品、聊天台词、站位或右键交互表现，避免气味、心理感受、镜头特写等 Minecraft 中难以确认的线索。",
        "NPC 如果即兴发挥，不能改变信物是否获得、阶段是否切换、体力是否消耗这些硬规则。",
        "当主角体力耗尽时，NPC 继续完成当前对话或场景收尾，再由阶段系统切换到下一场景。",
    ])

    doc.add_heading("十三 JSON 落地建议", level=1)
    add_p(doc, "每个角色每个阶段建议独立一个对话 JSON，文件名使用角色 ID 和阶段号，例如 student_scholar_phase2.json。小游戏可以放在对话节点中，也可以作为方块或物品交互 JSON。")
    add_bullets(doc, [
        "学生信物奖励建议配置为本轮物品，并在循环开始时清空。",
        "可保留信息建议使用单独道具或计分板标记，后续可在 JSON 条件里检测。",
        "消耗体力的回答用 staminaCost 控制，关键行动优先消耗 1 到 2 点。",
        "记忆翻牌不用手动写 gridSize，当前模组会规范为 15 格 3x5，每轮 6 个发光格子。",
        "节奏对抗建议用于知识、速度、调色等抽象竞争，不建议用作体力竞争。",
    ])

    doc.add_heading("十四 第一版制作优先级", level=1)
    add_table(doc, ["优先级", "内容", "目标"], [
        ["1", "完成 6 个阶段传送点和体力重置", "验证基础循环"],
        ["2", "做林澈、周燃、沈晴、许骁四条线", "覆盖四种小游戏"],
        ["3", "补唐玥、魏眠、何岚三条线", "形成七信物结构"],
        ["4", "加入老师最终三问", "完成真结局"],
        ["5", "补全失败锁线和跨循环保留信息", "形成多轮推理体验"],
    ], [2.0, 7.0, 7.0])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
