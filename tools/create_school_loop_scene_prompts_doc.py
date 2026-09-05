from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "学校循环场景提示第一版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    return p


def add_bullets(doc, items):
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "D9EAF7")
        if widths:
            header_cells[i].width = Cm(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("学校循环场景提示第一版")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("给 NPC 玩家自由发挥用的场景、目标与信息边界")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)

    add_para(doc, "定位：本文件不规定逐字台词。每个 NPC 玩家只需要按照场景目标、可透露信息和禁止越界内容进行即兴扮演。")
    add_para(doc, "规则底线：剧情阶段、体力消耗、信物发放、小游戏胜负和路线锁定仍由地图机制或主持指令决定，NPC 不靠口头承诺改变结果。")
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)

    add_cover(doc)

    add_heading(doc, "一、角色名单", 1)
    add_table(
        doc,
        ["角色", "角色ID建议", "定位", "主要玩法关系"],
        [
            ["主角", "protagonist", "唯一能推动全局剧情的人。其他玩家可以参与对话和小游戏，但不能推进阶段。", "收集信物、管理体力、选择路线。"],
            ["班长：BC", "monitor_bc", "秩序维护者，掌握出勤、值日、钥匙流转。", "记忆翻牌、巡查路线、纪律压力。"],
            ["体育生：狗头", "athlete_goutou", "冲动直接，习惯用比赛解决冲突。", "扳手腕、体育器材、力量路线。"],
            ["学霸：六谷", "scholar_liugu", "信息整理者，掌握题目、排名和资料。", "知识问答用普通对话实现。"],
            ["胆小鬼：随小乐", "timid_suixiaole", "看到过关键片段，但害怕承担后果。", "记忆翻牌、躲藏路线、目击证词。"],
            ["富家子：孙悟空", "rich_sunwukong", "资源多，习惯用交易和面子解决问题。", "抢柜子、道具交换、利益诱导。"],
            ["广播员：赵子龙", "broadcaster_zhaozilong", "控制广播室节奏，能把消息扩散给全校。", "节奏对抗、广播误导、时间提示。"],
            ["转学生：王少栋", "transfer_wangshaodong", "刚进入班级，身份和目的不透明。", "抢柜子、路线误导、隐藏物品。"],
            ["老师：雨哥", "teacher_yuge", "阶段控场者，负责把学校日程推进到下一个场景。", "全灭路线判定、阶段过渡、课堂/集合控制。"],
        ],
        [2.6, 3.2, 6.5, 5.0],
    )

    add_heading(doc, "二、NPC 即兴规则", 1)
    add_bullets(
        doc,
        [
            "NPC 可以自由发挥语气、态度、停顿、小动作和临场反应，但不能改变本场景的硬条件。",
            "NPC 可以暗示玩家去某个地点、找某个物品、参加某个小游戏；不能直接说出完整通关流程。",
            "信物不能保存到下一次循环。部分关键道具可以保留，但必须由地图或指令明确给予。",
            "不要使用无法在 Minecraft 内稳定验证的线索，例如气味、不可见心理状态、只有旁白知道的信息。",
            "生存和冒险模式玩家可以正常对话；创造模式或管理员负责选择 JSON、设置阶段、处理测试。",
        ],
    )

    add_heading(doc, "三、循环阶段", 1)
    add_para(doc, "每个阶段默认 5 点体力。体力耗尽后不立刻切阶段，等当前对话或小游戏结算结束，再进入下一阶段。")
    add_table(
        doc,
        ["阶段", "场景", "控场角色", "场景目标", "可自由发挥方向", "常见体力消耗"],
        [
            ["1", "校门口与早自习前", "雨哥、BC", "让主角认识所有人，确定第一条接触路线。", "催促集合、检查迟到、提醒今天有检查。", "帮人搬物、追问线索、进入小游戏。"],
            ["2", "教室与走廊", "BC、六谷、狗头", "引出纪律、成绩、体育器材三条路线。", "班长维持秩序，六谷关注题目，狗头挑衅比赛。", "扳手腕、翻找课桌、替人跑腿。"],
            ["3", "操场与器材室", "狗头、赵子龙", "让体力路线和广播路线发生交叉。", "比赛、广播通知、临时集合、器材丢失。", "扳手腕、节奏对抗、搬器材。"],
            ["4", "午休、柜子区与广播室", "孙悟空、王少栋、赵子龙", "处理资源交换、隐藏物品、广播误导。", "交易、借物、抢柜子、广播插播。", "抢柜子、追赶、交换条件。"],
            ["5", "实验室、办公室与放学前", "随小乐、雨哥", "把目击信息和老师路线推到明面。", "害怕被问责、办公室审问、临时检查。", "记忆翻牌、替人作证、销毁证据。"],
            ["6", "最终集合与循环重置", "雨哥、主角", "结算单人线或全收集线，准备进入下一轮。", "老师宣布结果，NPC 根据是否被解决表现不同态度。", "最后选择、补救行动、全灭判定。"],
        ],
        [1.2, 3.0, 2.5, 4.6, 5.0, 3.0],
    )

    add_heading(doc, "四、七条学生线", 1)
    add_table(
        doc,
        ["线路", "核心场景", "小游戏或机制", "NPC 场景提示", "成功后结果"],
        [
            ["BC 班长线", "教室点名、走廊巡查、值日表争夺。", "记忆翻牌：15 格中找 6 个发光格。", "BC 怀疑主角破坏纪律，可以逼主角解释行动，也可以用班规压人。", "获得班长信物，并解锁一条避开巡查的路线。"],
            ["狗头 体育生线", "操场挑战、器材室冲突。", "扳手腕：双方同步点击，进度条向对方方向推进，到头结束。", "狗头不轻易相信说服，倾向用胜负决定是否给线索。", "获得体育生信物，并能调动一次操场人流。"],
            ["六谷 学霸线", "早自习、考试资料、错题交换。", "知识问答：完全用普通对话选项实现，题目由作者自己写。", "六谷不需要说标准台词，只要围绕题目、资料和逻辑漏洞给压力。", "获得学霸信物，并解锁关键知识或密码。"],
            ["随小乐 胆小鬼线", "厕所门口、楼梯间、办公室外。", "记忆翻牌或普通对话压力选择。", "随小乐知道自己看到过什么，但害怕被牵连；可以说得含糊，让主角追问。", "获得胆小鬼信物，并补全某次事故的目击路线。"],
            ["孙悟空 富家子线", "柜子区、食堂、临时交易。", "抢柜子：点对加分，点错不扣分，一分钟内轮换目标。", "孙悟空可以谈条件、讲排场、用道具诱导主角改变路线。", "获得富家子信物，并获得可保留到循环外的资源型道具。"],
            ["赵子龙 广播员线", "广播室、操场集合、午休广播。", "节奏对抗：点击后立刻刷新下一条，目标位置和宽度随机。", "赵子龙通过广播制造全校误导，也可以临时帮主角转移注意力。", "获得广播员信物，并解锁一次全校通知效果。"],
            ["王少栋 转学生线", "校门口初见、柜子区、空教室。", "抢柜子或双人 UI 对抗。", "王少栋不熟悉学校规则，但知道某个不该知道的地点；说话可以显得谨慎。", "获得转学生信物，并揭示老师路线的前置条件。"],
        ],
        [2.2, 4.2, 3.8, 5.4, 4.2],
    )

    add_heading(doc, "五、老师雨哥线", 1)
    add_para(doc, "雨哥不是普通学生线。老师线的信物条件是：主角必须在单次循环内完成全部七名学生的事故路线，并在最终集合阶段触发老师结算。")
    add_bullets(
        doc,
        [
            "雨哥负责控场和过渡，例如上课、集合、检查、放学、进入办公室。",
            "雨哥可以追问主角为什么频繁离开，但不能直接封死所有路线，除非阶段或 JSON 已经设置失败条件。",
            "如果主角没有集齐七个学生信物，雨哥只给普通结局提示，不发放老师信物。",
            "如果主角在单次循环内集齐七个学生信物，雨哥进入最终审问或表彰反转场景，发放老师信物。",
        ],
    )

    add_heading(doc, "六、学霸六谷知识问答模板", 1)
    add_para(doc, "六谷的小游戏不用新 UI，直接用对话 JSON 的回答分支实现。题目、正确答案和惩罚由作者后续填写。")
    add_table(
        doc,
        ["节点", "场景作用", "设计方式"],
        [
            ["开场节点", "六谷确认主角是否真的理解线索。", "给出题目前置，不写固定语句，只写题目背景。"],
            ["题目节点", "提供 2 到 4 个回答。", "正确回答进入下一题或奖励节点，错误回答进入锁线节点。"],
            ["体力节点", "关键追问前扣体力。", "在回答前显示闪电图标，JSON 设置 staminaCost。"],
            ["成功节点", "给出学霸信物或关键密码。", "只给主角奖励，被互动者不获得。"],
            ["失败节点", "本轮不能继续学霸线。", "提示主角需要下次循环换顺序或带道具再来。"],
        ],
        [2.2, 5.0, 8.0],
    )

    add_heading(doc, "七、全收集路线写法", 1)
    add_bullets(
        doc,
        [
            "单条线可以独立完成，但全收集要求主角在同一轮内严格安排顺序。",
            "某些道具可以跨循环保留，用来降低下一轮体力消耗或提前打开捷径。",
            "信物只用于本轮结算，不跨循环保存；这样可以保证全收集路线仍然需要一次完整执行。",
            "建议把全收集路线拆成 6 个阶段，每阶段 5 点体力，并在每阶段安排 2 到 4 个消耗点。",
            "如果某一阶段体力使用错误，后面仍可完成单人线，但无法完成一轮全收集。",
        ],
    )

    add_heading(doc, "八、JSON 落地建议", 1)
    add_table(
        doc,
        ["内容", "建议字段", "说明"],
        [
            ["角色认领", "roleId", "使用固定角色ID，不绑定真实玩家名。"],
            ["阶段限制", "phase", "全服统一阶段；没有该阶段内容时，生存或冒险模式右键无效。"],
            ["自由场景", "scenePrompt", "写给 NPC 玩家看的场景目标，不当作固定台词显示。"],
            ["可透露信息", "allowedInfo", "NPC 可以自由表达这些信息。"],
            ["禁止内容", "forbiddenInfo", "避免 NPC 破坏路线或泄露结局。"],
            ["奖励", "rewards", "只给主角或右键发起者。"],
            ["体力", "staminaCost", "放在回答数据里，UI 用闪电图标提示，不把扣体力文字写进按钮。"],
            ["小游戏", "minigame", "可挂在对话选项，也可挂在右键物品或方块。"],
        ],
        [2.5, 3.3, 8.0],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
